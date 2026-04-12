from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "reports" / "bridge"
REPORT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "4472C4")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def create_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- TITLE PAGE ----
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Test Strategy & Test Plan")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("PlayReady RAG Chatbot - QA Automation Framework")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph("")

    info_table = [
        ("Document Version", "1.0"),
        ("Author", "Sushrut Nistane (KFORCE INC)"),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Project", "PlayReady Chatbot QA Automation"),
        ("Client", "Microsoft"),
        ("Status", "Draft"),
        ("Classification", "Confidential"),
    ]

    table = doc.add_table(rows=len(info_table), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(info_table):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        for p in table.rows[i].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
        set_cell_shading(table.rows[i].cells[0], "D6E4F0")

    doc.add_page_break()

    # ---- TABLE OF CONTENTS ----
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. Test Objectives",
        "4. Scope",
        "5. Test Strategy",
        "6. Test Planning",
        "7. Test Automation Architecture",
        "8. Evaluation Framework",
        "9. Reporting Strategy",
        "10. Compliance & Security Testing",
        "11. Risk Assessment",
        "12. Defect Management",
        "13. Approvals",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ---- 1. EXECUTIVE SUMMARY ----
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document defines the Test Strategy and Test Plan for the PlayReady RAG (Retrieval-Augmented Generation) "
        "Chatbot built on Azure AI Foundry. The chatbot serves as a knowledge assistant for PlayReady compliance, "
        "licensing, robustness rules, and technical documentation."
    )
    doc.add_paragraph(
        "The QA automation framework validates the chatbot's accuracy, safety, compliance awareness, and grounding "
        "against 42 PlayReady PDF documents using 965 scenario-based test cases. Evaluation is performed using both "
        "open-source (RAGAS) and Microsoft-native (Azure AI Evaluation SDK) metrics, covering quality, NLP accuracy, "
        "and AI safety dimensions."
    )
    doc.add_paragraph(
        "This framework is designed to be enterprise-grade, audit-safe, and compliant with Microsoft standards."
    )

    # ---- 2. PROJECT OVERVIEW ----
    doc.add_heading("2. Project Overview", level=1)
    add_styled_table(doc,
        ["Attribute", "Details"],
        [
            ["Application", "PlayReady RAG Chatbot"],
            ["Platform", "Azure AI Foundry"],
            ["Architecture", "Retrieval-Augmented Generation (RAG)"],
            ["Knowledge Base", "42 PlayReady PDF documents"],
            ["Total Test Cases", "965"],
            ["Evaluation Engines", "RAGAS + Azure AI Evaluation SDK"],
            ["Automation Stack", "Python, pytest, PyMuPDF, openpyxl"],
            ["CI/CD Ready", "Yes (pytest-based, command-line executable)"],
            ["Reporting", "Excel (color-coded) + HTML + JSON"],
        ],
    )

    # ---- 3. TEST OBJECTIVES ----
    doc.add_heading("3. Test Objectives", level=1)
    objectives = [
        "Validate that the chatbot provides accurate, grounded responses from PlayReady documentation.",
        "Ensure the chatbot correctly refuses out-of-scope, harmful, or sensitive queries.",
        "Verify compliance-related responses are accurate and complete per PlayReady rules.",
        "Detect hallucination — fabricated data, false agreements, or invented sections.",
        "Validate data security awareness — PII handling, encryption, access control, audit logging.",
        "Test conversational quality — tone adaptation, empathy, natural language handling.",
        "Ensure the chatbot cites correct documents and does not mix content across PDFs.",
        "Measure response quality using industry-standard NLP and AI-assisted metrics.",
        "Assess AI safety — violence, sexual content, self-harm, and hate/unfairness detection.",
        "Generate audit-ready reports with PASS/FAIL thresholds and color-coded Excel output.",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # ---- 4. SCOPE ----
    doc.add_heading("4. Scope", level=1)
    doc.add_heading("4.1 In-Scope", level=2)
    in_scope = [
        "Functional testing of chatbot responses against 42 PlayReady PDF documents.",
        "Scenario-based testing: positive, negative, edge case, hallucination, compliance, grounding.",
        "Data security and PII testing (5 scenarios per PDF = 210 test cases).",
        "Conversational and tone testing (casual, frustrated, formal, non-technical personas).",
        "Cross-document reasoning and version confusion testing.",
        "Prompt injection and jailbreak resistance testing.",
        "Competitor technology boundary testing (Widevine, FairPlay).",
        "Data leakage prevention testing.",
        "Automated evaluation using RAGAS (13 metrics) and Foundry SDK (12+ metrics).",
        "Excel reporting with thresholds, PASS/FAIL, and color coding.",
    ]
    for item in in_scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.2 Out-of-Scope", level=2)
    out_scope = [
        "UI/Frontend testing (Playwright-based UI E2E tests are separate).",
        "Performance and load testing of the Foundry agent.",
        "Infrastructure and deployment testing.",
        "Manual exploratory testing (covered separately).",
        "Testing of the Azure AI Foundry platform itself.",
    ]
    for item in out_scope:
        doc.add_paragraph(item, style="List Bullet")

    # ---- 5. TEST STRATEGY ----
    doc.add_heading("5. Test Strategy", level=1)

    doc.add_heading("5.1 Test Levels", level=2)
    add_styled_table(doc,
        ["Level", "Description", "Scope"],
        [
            ["Unit (Per-PDF)", "20 scenario-based tests per PDF document", "Within single document"],
            ["Integration (Cross-Doc)", "Tests spanning multiple PDF documents", "Across documents"],
            ["System (End-to-End)", "Full pipeline: PDF -> Test Generation -> Agent -> Evaluation -> Report", "Complete pipeline"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("5.2 Test Types & Distribution", level=2)
    add_styled_table(doc,
        ["#", "Test Type", "Per-PDF", "Global", "Total", "Severity", "Bug Type Detected"],
        [
            ["1", "Positive (Factual + Detail)", "84", "-", "84", "Medium", "Basic retrieval failures"],
            ["2", "Edge Case (Version/Negation/Multi-part)", "126", "-", "126", "High", "Boundary & logic errors"],
            ["3", "Negative (Out-of-scope/Wrong doc)", "84", "20", "104", "Critical", "Hallucination on unknown topics"],
            ["4", "Hallucination Traps", "84", "15", "99", "Critical", "Fabricated data & false agreements"],
            ["5", "Robustness (Paraphrase)", "42", "-", "42", "Medium", "Inconsistent answers"],
            ["6", "Compliance Rules", "42", "-", "42", "High", "Wrong compliance guidance"],
            ["7", "Grounding (Citation)", "42", "-", "42", "High", "Wrong document cited"],
            ["8", "Cross-Document", "42", "20", "62", "Critical", "Content blending across docs"],
            ["9", "Data Security / PII", "210", "-", "210", "Critical", "PII exposure, weak security guidance"],
            ["10", "Conversational / Tone", "84", "25", "109", "Medium", "Robotic/inappropriate tone"],
            ["11", "Version Confusion", "-", "10", "10", "High", "Outdated rules applied"],
            ["12", "False Premise", "-", "15", "15", "Critical", "Agent agrees with false claims"],
            ["13", "Competitor Technology", "-", "10", "10", "High", "Scope boundary violations"],
            ["14", "Data Leakage", "-", "10", "10", "Critical", "Sensitive info exposure"],
            ["", "TOTAL", "840", "125", "965", "", ""],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("5.3 Test Data Strategy", level=2)
    add_styled_table(doc,
        ["Attribute", "Details"],
        [
            ["Knowledge Base", "42 PlayReady PDF documents in data/kb/"],
            ["PDF Registry", "data/pdf_registry.json — tracks doc_id, version, category, active status"],
            ["Chunk Registry", "data/chunk_registry.json — tracks chunk_id, pages, hash, text preview"],
            ["Chunk Size", "1000 characters with 200-character overlap"],
            ["Ground Truth", "Auto-extracted from PDF chunks (first 500 chars)"],
            ["Test Cases File", "data/test_cases.json (965 cases, pipeline input)"],
            ["Backup Strategy", "Auto-backup before regeneration (timestamped)"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("5.4 Tools & Technologies", level=2)
    add_styled_table(doc,
        ["Tool", "Purpose", "Version"],
        [
            ["Python", "Core automation language", "3.14"],
            ["pytest", "Test execution framework", "9.0+"],
            ["Azure AI Foundry SDK", "Agent interaction (RAG queries)", "Latest"],
            ["RAGAS", "Open-source RAG evaluation (13 metrics)", "Latest"],
            ["azure-ai-evaluation", "Microsoft AI evaluation SDK (12+ metrics)", "1.0+"],
            ["PyMuPDF (fitz)", "PDF text extraction and chunking", "Latest"],
            ["openpyxl", "Excel report generation with formatting", "Latest"],
            ["pandas", "Data processing and analysis", "Latest"],
            ["VS Code", "Development IDE", "Latest"],
            ["Azure DevOps (ADO)", "Defect tracking and CI/CD", "Cloud"],
        ],
    )

    # ---- 6. TEST PLANNING ----
    doc.add_heading("6. Test Planning", level=1)

    doc.add_heading("6.1 Test Environment", level=2)
    add_styled_table(doc,
        ["Component", "Details"],
        [
            ["OS", "Windows 11"],
            ["Python", "3.14.3"],
            ["IDE", "VS Code"],
            ["Agent Platform", "Azure AI Foundry"],
            ["Authentication", "Azure OpenAI API Key + Azure Identity"],
            ["Network", "Microsoft corporate network (VPN)"],
            ["Config", ".env file (endpoint, API key, deployment, project settings)"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("6.2 Entry Criteria", level=2)
    entry = [
        "All 42 PlayReady PDFs are uploaded to data/kb/ with clean filenames.",
        "pdf_registry.json is generated and verified.",
        "Azure AI Foundry agent is deployed and accessible via API.",
        ".env file contains valid Azure OpenAI credentials.",
        "All Python dependencies are installed in virtual environment.",
        "Test case generation scripts execute without errors.",
        "100% scenario coverage validated (0 gaps in coverage report).",
    ]
    for item in entry:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.3 Exit Criteria", level=2)
    exit_c = [
        "All 965 test cases executed against the Foundry agent.",
        "RAGAS evaluation completed with scores for all applicable metrics.",
        "Foundry evaluation completed (Quality + NLP + Safety).",
        "All critical severity test failures documented as ADO bugs.",
        "Coverage report shows 100% scenario coverage.",
        "All Excel reports generated with PASS/FAIL and color coding.",
        "No critical/blocker defects remain unresolved.",
        "Test Strategy and Planning document reviewed and approved.",
    ]
    for item in exit_c:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.4 Test Execution Flow", level=2)
    flow_steps = [
        "Phase 1: PDF Ingestion — Upload PDFs, rename, generate registry, extract chunks.",
        "Phase 2: Test Generation — Run 3 generator scripts (scenario + negative + conversational).",
        "Phase 3: Coverage Validation — Run validate_test_coverage.py, confirm 100% coverage.",
        "Phase 4: Agent Execution — Run query_foundry_agent.py to call Foundry agent with all 965 questions.",
        "Phase 5: RAGAS Evaluation — Run pytest -m ragas to evaluate using RAGAS metrics.",
        "Phase 6: Foundry Evaluation — Run pytest -m foundry_eval to evaluate using Foundry SDK metrics.",
        "Phase 7: Reporting — Open Excel reports, review PASS/FAIL, identify bugs.",
        "Phase 8: Bug Filing — Create ADO bugs for failed test cases.",
        "Phase 9: Retest — Fix issues, regenerate, re-evaluate.",
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("6.5 Test Schedule", level=2)
    add_styled_table(doc,
        ["Phase", "Activity", "Duration", "Status"],
        [
            ["Phase 1", "PDF ingestion, cleanup, registry generation", "1 day", "Completed"],
            ["Phase 2", "Test case generation (965 cases)", "1 day", "Completed"],
            ["Phase 3", "Coverage validation (100%)", "0.5 day", "Completed"],
            ["Phase 4", "Foundry agent execution (965 queries)", "1 day", "Pending"],
            ["Phase 5", "RAGAS evaluation", "0.5 day", "Pending"],
            ["Phase 6", "Foundry SDK evaluation", "0.5 day", "Pending"],
            ["Phase 7", "Report review and bug filing", "1 day", "Pending"],
            ["Phase 8", "Retest after fixes", "1-2 days", "Pending"],
        ],
    )

    # ---- 7. TEST AUTOMATION ARCHITECTURE ----
    doc.add_heading("7. Test Automation Architecture", level=1)

    doc.add_heading("7.1 Project Structure", level=2)
    structure = """playready-qa-automation/
├── data/
│   ├── kb/                              (42 PlayReady PDFs)
│   ├── pdf_registry.json                (PDF metadata tracking)
│   ├── chunk_registry.json              (Chunk metadata with IDs, pages, hashes)
│   ├── test_cases.json                  (965 test cases - pipeline input)
│   ├── test_cases_master.json           (Per-PDF backup)
│   ├── test_cases_negative_master.json  (Global negative backup)
│   ├── test_cases_conversational_master.json (Global conversational backup)
│   └── ragas_eval_dataset.json          (Agent responses - evaluation input)
├── scripts/
│   ├── generate_ragas_testset.py        (Per-PDF test generator - 20 per PDF)
│   ├── generate_negative_testcases.py   (Global negative test generator)
│   ├── generate_conversational_testcases.py (Global conversational test generator)
│   ├── validate_test_coverage.py        (100% coverage validator)
│   ├── query_foundry_agent.py           (Foundry agent caller)
│   └── rename_pdfs.py                   (PDF filename cleanup)
├── ragas_layer/
│   └── ragas_runner.py                  (RAGAS evaluation engine)
├── foundry_layer/
│   └── foundry_evaluator.py             (Foundry SDK evaluation engine)
├── tests/
│   ├── test_ragas_eval.py               (RAGAS pytest trigger)
│   └── test_foundry_eval.py             (Foundry pytest trigger)
├── audit/
│   └── reporting.py                     (RAGAS report generator)
├── reports/bridge/
│   ├── Test_Coverage_Report.xlsx
│   ├── Bridge_Evaluation_Report.xlsx    (RAGAS)
│   └── Foundry_Evaluation_Report.xlsx
└── artifacts/
    ├── ragas/                           (RAGAS JSON/CSV results)
    └── foundry/                         (Foundry JSON/CSV results)"""

    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    doc.add_heading("7.2 Key Files & Responsibilities", level=2)
    add_styled_table(doc,
        ["File", "Purpose", "Input", "Output"],
        [
            ["generate_ragas_testset.py", "Generate 20 per-PDF test cases", "PDFs + registry", "test_cases.json (840)"],
            ["generate_negative_testcases.py", "Generate global negative tests", "registry", "Merges 100 into test_cases.json"],
            ["generate_conversational_testcases.py", "Generate conversational tests", "None", "Merges 25 into test_cases.json"],
            ["validate_test_coverage.py", "Validate 100% scenario coverage", "test_cases.json", "Coverage Excel + JSON"],
            ["query_foundry_agent.py", "Call Foundry agent with test cases", "test_cases.json", "ragas_eval_dataset.json"],
            ["ragas_runner.py", "Run RAGAS evaluation metrics", "ragas_eval_dataset.json", "Excel + JSON + CSV"],
            ["foundry_evaluator.py", "Run Foundry SDK evaluation metrics", "ragas_eval_dataset.json", "Excel + JSON + CSV"],
        ],
    )

    # ---- 8. EVALUATION FRAMEWORK ----
    doc.add_heading("8. Evaluation Framework", level=1)

    doc.add_heading("8.1 RAGAS Metrics", level=2)
    add_styled_table(doc,
        ["Metric", "What It Measures", "Scale", "Direction"],
        [
            ["Faithfulness", "Is the answer faithful to the context?", "0-1", "Higher is better"],
            ["Answer Relevancy", "Is the answer relevant to the question?", "0-1", "Higher is better"],
            ["Context Precision", "Are retrieved contexts precise?", "0-1", "Higher is better"],
            ["Context Recall", "Are all relevant contexts retrieved?", "0-1", "Higher is better"],
            ["Answer Correctness", "Is the answer factually correct?", "0-1", "Higher is better"],
            ["Answer Similarity", "Semantic similarity to ground truth", "0-1", "Higher is better"],
            ["Harmfulness", "Does the answer contain harmful content?", "0-1", "Lower is better"],
            ["Maliciousness", "Is the answer malicious?", "0-1", "Lower is better"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("8.2 Foundry Quality Metrics (AI-Assisted, Scale 1-5)", level=2)
    add_styled_table(doc,
        ["Metric", "What It Measures", "Threshold", "Pass Criteria"],
        [
            ["Coherence", "Logical coherence of response", ">= 3.0", "Score >= 3.0"],
            ["Fluency", "Grammatical correctness", ">= 3.0", "Score >= 3.0"],
            ["Relevance", "Relevance to query", ">= 3.0", "Score >= 3.0"],
            ["Groundedness", "Grounded in provided context", ">= 3.0", "Score >= 3.0"],
            ["Similarity", "Similarity to ground truth", ">= 3.0", "Score >= 3.0"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("8.3 Foundry NLP Metrics (Math-Based, Scale 0-1)", level=2)
    add_styled_table(doc,
        ["Metric", "What It Measures", "Threshold", "Requires LLM"],
        [
            ["F1 Score", "Token-level precision + recall", ">= 0.3", "No"],
            ["ROUGE Score", "Recall-oriented overlap", ">= 0.3", "No"],
            ["BLEU Score", "Precision-oriented n-gram overlap", ">= 0.2", "No"],
            ["METEOR Score", "Semantic matching", ">= 0.3", "No"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("8.4 Foundry Safety Metrics (Scale 0-7, Lower = Safer)", level=2)
    add_styled_table(doc,
        ["Metric", "What It Measures", "Threshold", "Pass Criteria"],
        [
            ["Violence", "Detects violent content", "<= 1", "Score <= 1"],
            ["Sexual", "Detects sexual content", "<= 1", "Score <= 1"],
            ["Self-Harm", "Detects self-harm content", "<= 1", "Score <= 1"],
            ["Hate/Unfairness", "Detects hate speech and bias", "<= 1", "Score <= 1"],
        ],
    )

    # ---- 9. REPORTING STRATEGY ----
    doc.add_heading("9. Reporting Strategy", level=1)
    add_styled_table(doc,
        ["Report", "Content", "Format", "Color Coding"],
        [
            ["Test_Coverage_Report.xlsx", "100% scenario coverage validation per PDF", "Excel (6 tabs)", "Green=COMPLETE, Red=INCOMPLETE"],
            ["Bridge_Evaluation_Report.xlsx", "RAGAS metrics with thresholds and per-row scores", "Excel (4 tabs)", "Green=PASS, Red=FAIL, Yellow=SKIPPED"],
            ["Foundry_Evaluation_Report.xlsx", "Foundry Quality + NLP + Safety metrics", "Excel (6 tabs)", "Green=PASS, Red=FAIL, Yellow=SKIPPED"],
            ["pytest_report.html", "Detailed test execution log", "HTML", "Standard pytest HTML"],
            ["coverage_report.json", "Machine-readable coverage data", "JSON", "N/A"],
        ],
    )

    # ---- 10. COMPLIANCE & SECURITY TESTING ----
    doc.add_heading("10. Compliance & Security Testing", level=1)
    doc.add_paragraph(
        "Given that this is a Microsoft client project, data security and compliance testing is a top priority. "
        "The framework includes 210 dedicated data security test cases (5 per PDF) covering:"
    )
    add_styled_table(doc,
        ["Security Area", "Test Count", "What It Validates"],
        [
            ["PII Exposure", "42", "Agent handles personally identifiable information correctly"],
            ["Data Retention", "42", "Agent provides accurate data retention/deletion guidance"],
            ["Encryption Keys", "42", "Agent gives secure key management advice"],
            ["Access Control", "42", "Agent recommends proper authorization (not overly permissive)"],
            ["Audit Logging", "42", "Agent covers mandatory audit logging requirements"],
            ["Prompt Injection", "10 (global)", "Agent resists jailbreak and manipulation attempts"],
            ["Data Leakage", "10 (global)", "Agent does not expose internal/confidential information"],
        ],
    )

    # ---- 11. RISK ASSESSMENT ----
    doc.add_heading("11. Risk Assessment", level=1)
    add_styled_table(doc,
        ["Risk", "Impact", "Probability", "Mitigation"],
        [
            ["Foundry agent downtime", "High", "Low", "Retry logic with delays; run during off-peak hours"],
            ["API rate limiting", "Medium", "Medium", "2-second delay between calls; batch execution"],
            ["PDF text extraction failures", "Medium", "Low", "PyMuPDF fallback; manual review of failed PDFs"],
            ["Ground truth inaccuracy", "High", "Medium", "Auto-extracted from chunks; manual review for compliance"],
            ["Evaluation metric instability", "Medium", "Low", "Run both RAGAS and Foundry for cross-validation"],
            ["Knowledge base updates", "High", "Medium", "Version tracking via pdf_registry.json and chunk_registry.json"],
            ["Test data staleness", "Medium", "Medium", "Auto-backup and regeneration workflow"],
            ["False positive/negative in safety", "Medium", "Low", "Manual review of safety metric edge cases"],
        ],
    )

    # ---- 12. DEFECT MANAGEMENT ----
    doc.add_heading("12. Defect Management", level=1)

    doc.add_heading("12.1 Severity Levels", level=2)
    add_styled_table(doc,
        ["Severity", "Definition", "Examples", "SLA"],
        [
            ["Critical (P0)", "Agent produces harmful, misleading, or security-violating output", "Hallucination, PII exposure, prompt injection bypass", "Immediate fix"],
            ["High (P1)", "Agent provides incorrect or incomplete information", "Wrong compliance rules, version confusion, wrong citation", "Fix within 2 days"],
            ["Medium (P2)", "Agent response quality is suboptimal", "Poor tone, overly verbose, partial answer", "Fix within 5 days"],
            ["Low (P3)", "Minor cosmetic or formatting issues", "Inconsistent formatting, minor wording issues", "Next sprint"],
        ],
    )

    doc.add_heading("12.2 Bug Tracking", level=2)
    doc.add_paragraph("All defects are tracked in Azure DevOps (ADO) with the following workflow:")
    workflow = [
        "Bug identified from evaluation report (PASS/FAIL analysis).",
        "ADO bug created with: Title, Repro Steps, Expected vs Actual, Severity, Test Case ID.",
        "Assigned to development team for investigation.",
        "Fix implemented and deployed to Foundry agent.",
        "Retest using the same test case ID — re-run specific category if needed.",
        "Bug closed with evidence (updated evaluation report showing PASS).",
    ]
    for step in workflow:
        doc.add_paragraph(step, style="List Number")

    # ---- 13. APPROVALS ----
    doc.add_heading("13. Approvals", level=1)
    add_styled_table(doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["QA Engineer", "Sushrut Nistane", "", ""],
            ["QA Lead", "", "", ""],
            ["Project Manager", "", "", ""],
            ["Client Stakeholder", "", "", ""],
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- End of Document ---")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    output_path = REPORT_DIR / "Test_Strategy_and_Planning.docx"
    doc.save(str(output_path))
    print(f"Document saved: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    create_document()